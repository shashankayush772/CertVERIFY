from __future__ import annotations

from datetime import date
from typing import Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_TITLE = (
    "Design and Development of a Tamper-Proof Academic Credential Verification System "
    "Using Cryptographic Integrity and AI-Based Forgery Detection"
)

OUTPUT_PATH = r"C:\Users\Silpa\certverify\CertVerify_Project_Report_Phase1.docx"


def _set_run_font(run, *, name: str = "Times New Roman", size_pt: int = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size_pt: int = 12,
    align: Optional[int] = None,
    space_after_pt: int = 0,
    space_before_pt: int = 0,
    keep_together: bool = False,
    keep_with_next: bool = False,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 2  # double spacing
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)
    pf.keep_together = keep_together
    pf.keep_with_next = keep_with_next

    run = p.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold)
    return p


def _add_chapter_heading(doc: Document, text: str):
    return _add_paragraph(
        doc,
        text,
        bold=True,
        size_pt=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=6,
        keep_with_next=True,
    )


def _add_section_heading(doc: Document, text: str):
    return _add_paragraph(doc, text, bold=True, size_pt=12, space_after_pt=3, keep_with_next=True)


def _add_page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_run_font(run, size_pt=12, bold=False)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _set_document_defaults(doc: Document):
    # Page setup: margins as specified
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)

    # Header with project title
    header = section.header
    header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_p.clear()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run(PROJECT_TITLE)
    _set_run_font(header_run, size_pt=12, bold=True)

    # Footer with centered page numbers
    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.clear()
    _add_page_number_field(footer_p)

    # Normal style: Times New Roman 12, double spacing
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)


def _add_table(doc: Document, rows: int, cols: int, *, style: str = "Table Grid"):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    return table


def _set_cell_text(cell, text: str, *, bold: bool = False):
    # Ensure a single paragraph with our font settings
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 2
    run = p.add_run(text)
    _set_run_font(run, bold=bold)


def build_report() -> Tuple[str, List[str], List[Tuple[str, str]]]:
    doc = Document()
    _set_document_defaults(doc)

    student_fill_markers: List[str] = []
    figures: List[Tuple[str, str]] = []

    # -------------------------
    # Front matter (minimal)
    # -------------------------
    _add_paragraph(doc, "CERTVERIFY", bold=True, size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    _add_paragraph(doc, PROJECT_TITLE, bold=True, size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=12)

    _add_paragraph(
        doc,
        "Department of Information Science and Engineering\n"
        "New Horizon College of Engineering, Bengaluru\n"
        "Affiliated to Visvesvaraya Technological University (VTU), Belagavi",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=12,
    )

    student_fill_markers.extend(
        [
            "[STUDENT TO FILL: Student Name(s)]",
            "[STUDENT TO FILL: USN(s)]",
            "[STUDENT TO FILL: Semester/Section]",
            "[STUDENT TO FILL: Academic Year]",
            "[STUDENT TO FILL: Guide Name and Designation]",
            "[STUDENT TO FILL: HOD Name and Designation]",
            "[STUDENT TO FILL: Principal/Director Name (if required by template)]",
            "[STUDENT TO FILL: Place and Date of Submission]",
        ]
    )

    _add_paragraph(doc, f"[STUDENT TO FILL: Student Name(s)]", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"[STUDENT TO FILL: USN(s)]", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"[STUDENT TO FILL: Semester/Section]", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"[STUDENT TO FILL: Academic Year]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=12)
    _add_paragraph(doc, f"Guide: [STUDENT TO FILL: Guide Name and Designation]", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"HOD: [STUDENT TO FILL: HOD Name and Designation]", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"Submitted on: [STUDENT TO FILL: Place and Date of Submission]", align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # -------------------------
    # SECTION 1: Abstract (<=250 words)
    # -------------------------
    _add_section_heading(doc, "SECTION 1: ABSTRACT")
    abstract = (
        "Academic credential fraud and post-issuance tampering have become increasingly prevalent due to the ease of "
        "high-quality digital editing, leading to verification delays, reputational damage to institutions, and "
        "significant risk for employers and universities. This project presents CertVerify, a web-based academic "
        "credential issuance and verification system that combines cryptographic integrity and AI-based forgery "
        "detection to deliver defense-in-depth validation. During issuance, institutions upload certificate files "
        "(PDF/JPG/PNG), after which the system computes a SHA-256 hash, generates an RSA-2048 digital signature, and "
        "produces a QR code for public verification; the hash is also anchored on the Polygon Amoy testnet via a "
        "Solidity smart contract to provide immutable auditability. During verification, four sequential layers are "
        "executed: (i) visual forgery detection using an ELA-preprocessed ResNet-18 CNN trained on CASIA v2.0, "
        "(ii) stored versus recomputed SHA-256 hash integrity comparison, (iii) RSA signature validation using the "
        "issuer’s public key, and (iv) blockchain hash cross-check for tamper evidence and lifecycle status. The "
        "forgery detection model, fine-tuned for 15 epochs with Adam (\u03b7=0.0005) and CosineAnnealingLR on 12,614 "
        "CASIA images, achieved 94.67% accuracy with Precision 0.9564, Recall 0.9360, and F1-score 0.9461. The "
        "system outputs AUTHENTIC, TAMPERED, FORGED, or INVALID with confidence scores, enabling fast, reliable, and "
        "auditable credential verification suitable for institutional adoption and extensible for future deployment."
    )
    _add_paragraph(doc, abstract)

    doc.add_page_break()

    # -------------------------
    # SECTION 2: Acknowledgement
    # -------------------------
    _add_section_heading(doc, "SECTION 2: ACKNOWLEDGEMENT")
    ack = (
        "We express our sincere gratitude to the Department of Information Science and Engineering, New Horizon "
        "College of Engineering, Bengaluru, for providing the opportunity and resources to carry out this project. "
        "We convey our profound thanks to our project guide, [STUDENT TO FILL: Guide Name and Designation], for "
        "continuous guidance, technical suggestions, and encouragement throughout the project work. We are thankful "
        "to [STUDENT TO FILL: HOD Name and Designation], Head of the Department, for valuable support and for "
        "facilitating the necessary facilities. We also thank the Principal/Management of New Horizon College of "
        "Engineering for creating a motivating academic environment. Finally, we extend our gratitude to all "
        "faculty members, laboratory staff, and our friends for their cooperation, and to our parents for their "
        "constant support and motivation during the successful completion of this work."
    )
    _add_paragraph(doc, ack)

    doc.add_page_break()

    # -------------------------
    # SECTION 3: Table of Contents (approx page numbers)
    # -------------------------
    _add_section_heading(doc, "SECTION 3: TABLE OF CONTENTS")
    toc_lines = [
        "SECTION 1: Abstract .......................................................... 1",
        "SECTION 2: Acknowledgement .................................................... 2",
        "SECTION 3: Table of Contents .................................................. 3",
        "CHAPTER 1: Introduction ....................................................... 4",
        "  1.1 Domain Introduction ...................................................... 4",
        "  1.2 Problem Statement ........................................................ 6",
        "  1.3 Existing System .......................................................... 7",
        "  1.4 Proposed System .......................................................... 8",
        "CHAPTER 2: Literature Survey .................................................. 10",
        "CHAPTER 3: Analysis of Reviewed Papers ........................................ 20",
        "CHAPTER 4: Requirement Analysis ............................................... 23",
        "  4.1 Functional Requirements .................................................. 23",
        "  4.2 Non-Functional Requirements .............................................. 25",
        "  4.3 Hardware Requirements .................................................... 26",
        "  4.4 Software Requirements .................................................... 27",
        "CHAPTER 5: Design ............................................................. 29",
        "  5.1 System Architecture ...................................................... 29",
        "  5.2 Data Flow Diagram ........................................................ 31",
        "  5.3 Algorithm ............................................................... 33",
        "  5.4 Database Design .......................................................... 36",
        "CHAPTER 6: Conclusion .......................................................... 40",
        "REFERENCES .................................................................... 42",
    ]
    for line in toc_lines:
        _add_paragraph(doc, line)

    doc.add_page_break()

    # -------------------------
    # CHAPTER 1 — Introduction
    # -------------------------
    _add_chapter_heading(doc, "CHAPTER 1 — INTRODUCTION")

    _add_section_heading(doc, "1.1 Domain Introduction")
    intro_paras = [
        (
            "Academic credentials such as degree certificates, transcripts, and provisional letters form the backbone of "
            "trust in higher education and employment ecosystems. Institutions are expected to guarantee that a certificate "
            "represents an authentic academic outcome, while recruiters and universities rely on these documents to make "
            "high-stakes decisions. However, the digitization of workflows and widespread availability of editing tools have "
            "reduced the cost of document manipulation, increasing both the volume and sophistication of credential fraud."
        ),
        (
            "Credential fraud in practice appears in multiple forms: completely fabricated certificates, visually altered "
            "documents where names/grades/registration numbers are edited, and post-issuance tampering where an originally "
            "valid certificate is modified after being legitimately issued. Traditional verification approaches—manual checking "
            "with issuing institutions, stamped photocopies, and email confirmations—are slow, non-scalable, and prone to "
            "human error, leading to operational bottlenecks for registrars and verification cells."
        ),
        (
            "Modern digital approaches have attempted to address these issues through cryptographic hashing, digital signatures, "
            "and increasingly, blockchain-based registries. Hash-based verification can detect bit-level changes, while digital "
            "signatures establish issuer authenticity and non-repudiation. Blockchain solutions promise immutability and public "
            "auditability. Nevertheless, systems that rely exclusively on these mechanisms can still fail when the input itself "
            "is a convincing visual forgery that was never officially issued—particularly when a forged certificate is produced "
            "from scratch or derived from templates without any prior reference record."
        ),
        (
            "Therefore, a robust academic credential verification system requires both cryptographic integrity controls and "
            "content-level forgery detection. Advances in deep learning-based image forensics enable the classification of "
            "tampered versus authentic images by learning artifact patterns that are difficult to perceive manually. When combined "
            "with cryptographic hashing, RSA signatures, and on-chain anchoring for audit, a layered design can deliver high "
            "assurance and practical usability. CertVerify is developed in this context to provide a deployable, web-based "
            "system that supports institutional issuance, public verification, and comprehensive tamper/forgery detection."
        ),
    ]
    for p in intro_paras:
        _add_paragraph(doc, p)

    _add_section_heading(doc, "1.2 Problem Statement")
    _add_paragraph(
        doc,
        "Existing academic credential verification mechanisms suffer from multiple practical and security limitations, "
        "which motivate the need for an integrated, tamper-proof and forgery-aware solution. The key issues are:",
    )
    disadvantages = [
        "Manual verification is time-consuming and does not scale with large volumes of requests.",
        "Email/phone-based confirmation lacks standardization and is vulnerable to social engineering.",
        "Hash-only systems can detect file tampering but cannot identify visually forged documents created without official issuance.",
        "Blockchain-only registries may prove immutability, but do not evaluate whether the certificate content itself is forged.",
        "Many existing solutions omit revocation and lifecycle management, making it difficult to handle withdrawals, corrections, or re-issuance.",
        "Users lack a transparent, layer-wise explanation of verification outcomes and confidence, reducing trust in automated decisions.",
    ]
    for d in disadvantages:
        _add_paragraph(doc, f"• {d}")

    _add_section_heading(doc, "1.3 Existing System")
    existing = [
        (
            "Manual systems: In many institutions, verification is performed by checking physical registers or internal student "
            "information systems and responding to queries via email or official letters. While authoritative, these methods require "
            "human involvement for every request, resulting in delays, inconsistent formats, and potential transcription errors. "
            "Additionally, manual checks provide limited audit trails and are difficult to integrate into automated recruitment workflows."
        ),
        (
            "Hash-based systems: Some solutions generate a cryptographic hash (e.g., SHA-256) of the certificate file and later verify "
            "integrity by recomputing and comparing the hash. This approach is effective for detecting bit-level changes to a known issued "
            "document, but it assumes that the verifier has access to the correct stored hash and that the certificate was originally issued "
            "by the system. A visually forged certificate created independently may still pass superficial inspection even though no official "
            "record exists."
        ),
        (
            "Blockchain-only systems: Blockchain registries store certificate identifiers and hashes on-chain to provide immutability and public "
            "auditability. However, such solutions typically do not incorporate content-level analysis and can only assert that a certain hash "
            "was anchored at a certain time. They often introduce operational complexity (transaction fees, latency, dependency on RPC providers) "
            "and may not include revocation mechanisms or institutional lifecycle workflows. As a result, they cannot fully address the spectrum "
            "of academic document fraud."
        ),
    ]
    for p in existing:
        _add_paragraph(doc, p)

    _add_section_heading(doc, "1.4 Proposed System")
    proposed = (
        "CertVerify proposes a defense-in-depth academic credential verification system that integrates four complementary verification "
        "layers and supports two stakeholder roles: Institution and Verifier. During certificate issuance, an authenticated institution user "
        "uploads the certificate (PDF/JPG/PNG). The system generates a SHA-256 hash, produces an RSA-2048 digital signature using institution "
        "keys, stores certificate metadata in a relational database, and generates a QR code that encodes the certificate identifier and "
        "verification endpoint. For external immutability, the certificate hash is also anchored on the Polygon blockchain (Amoy testnet) via a "
        "Solidity smart contract, enabling third-party audit and tamper evidence."
    )
    _add_paragraph(doc, proposed)
    proposed2 = (
        "During verification, CertVerify executes a four-layer sequential pipeline with early-exit decisions and an auditable report: "
        "(1) AI-based forgery detection using ELA preprocessing followed by ResNet-18 inference to estimate whether the certificate is visually "
        "tampered or fabricated; (2) SHA-256 integrity verification to detect bit-level modifications of an issued file; (3) RSA-2048 signature "
        "validation to confirm issuer authenticity and non-repudiation using stored public key material; and (4) Polygon blockchain cross-check "
        "to confirm the stored hash is consistent with the on-chain record and has not been revoked. The system returns one of four outcomes—"
        "AUTHENTIC, TAMPERED, FORGED, or INVALID—along with confidence scores and layer-wise results, thereby improving reliability, "
        "transparency, and scalability compared to existing approaches."
    )
    _add_paragraph(doc, proposed2)

    doc.add_page_break()

    # -------------------------
    # CHAPTER 2 — Literature Survey (10 papers)
    # -------------------------
    _add_chapter_heading(doc, "CHAPTER 2 — LITERATURE SURVEY")

    def paper_block(
        title: str,
        authors_year: str,
        did: str,
        methodology: str,
        advantages: List[str],
        disadvantages: List[str],
        how_addressed: str,
    ):
        _add_section_heading(doc, title)
        _add_paragraph(doc, f"Authors and Year: {authors_year}")
        _add_paragraph(doc, f"What they did: {did}")
        _add_paragraph(doc, f"Methodology: {methodology}")
        _add_paragraph(doc, "Advantages:")
        for a in advantages:
            _add_paragraph(doc, f"• {a}")
        _add_paragraph(doc, "Disadvantages / Limitations:")
        for d in disadvantages:
            _add_paragraph(doc, f"• {d}")
        _add_paragraph(doc, f"How CertVerify addresses the limitations: {how_addressed}")

    # 1
    paper_block(
        "2.1 Golar et al. (2025) — DOI: 10.65521/ijacect.v14i3s.1629",
        "Golar et al., 2025",
        "They proposed a combined framework for academic credential verification using cryptographic techniques and distributed ledgers to reduce certificate fraud.",
        "Conceptual architecture describing certificate issuance, hashing, blockchain storage, and basic verification steps, primarily positioned as a design proposal.",
        [
            "Recognizes the need for combining multiple controls rather than relying on a single mechanism.",
            "Highlights blockchain immutability as a means of strengthening auditability.",
        ],
        [
            "Largely conceptual; lacks an implemented end-to-end system and deployed evaluation.",
            "Does not include AI-based visual forgery detection; forged certificates not issued by the system remain a gap.",
            "Limited discussion of operational workflows such as user roles, revocation, and verification logging.",
        ],
        "CertVerify implements a complete working system with institution/verifier roles, issuance workflow, QR-based public verification, audit logs, revocation support, and an AI forgery detection layer to handle visually forged documents.",
    )

    # 2
    paper_block(
        "2.2 Ch et al. (2024) — DOI: 10.6977/IJoSI.202403_8(1).0001",
        "Ch et al., 2024",
        "They explored deep learning for document forgery detection under constrained training data, focusing on classification of authentic vs manipulated documents.",
        "Deep learning-based classifier trained on a limited dataset; emphasis on feature learning for forgery cues and evaluation on the given dataset split.",
        [
            "Demonstrates the effectiveness of CNN-based approaches for forgery detection.",
            "Shows potential for automation in document authenticity assessment.",
        ],
        [
            "Dataset scale is limited, reducing generalization and comparability across benchmarks.",
            "Does not integrate cryptographic integrity checks or issuer authentication mechanisms.",
        ],
        "CertVerify uses CASIA v2.0 as a widely recognized benchmark (12,614 images) with a clear training protocol (ELA + ResNet-18 fine-tuning) and combines AI with SHA-256 hashing and RSA signatures for complete end-to-end verification.",
    )

    # 3
    paper_block(
        "2.3 Babu et al. (2022) — IJACSA 13(5)",
        "Babu et al., 2022",
        "They presented a certificate verification framework centered on cryptographic hashing and/or blockchain mechanisms to validate integrity of issued credentials.",
        "Certificate records are hashed and compared during verification; in some variants, hashes are stored in a distributed ledger to prevent modification of stored records.",
        [
            "Strong at detecting post-issuance tampering when the original hash is securely stored.",
            "Supports automation and reduces manual verification overhead.",
        ],
        [
            "Hash-only verification cannot detect visually forged certificates created independently.",
            "May not provide issuer non-repudiation unless signatures and key management are incorporated.",
            "Often lacks interpretability of outcomes beyond match/mismatch.",
        ],
        "CertVerify adds AI-based visual forgery detection as Layer 1 to detect non-issued forgeries, and uses RSA-2048 signatures to establish issuer authenticity, along with blockchain anchoring for auditability.",
    )

    # 4
    paper_block(
        "2.4 Vidal et al. (2020) — DOI: 10.1109/ICSESS49925.2020.9237735",
        "Vidal et al., 2020",
        "They investigated blockchain-based academic certificate management and discussed integrity, transparency, and operational issues such as revocation and lifecycle management.",
        "Blockchain-backed certificate registry with focus on integrity and traceability; analysis of challenges around revocation, updates, and governance.",
        [
            "Identifies practical lifecycle concerns that many blockchain prototypes ignore.",
            "Emphasizes the importance of governance and revocation for real deployments.",
        ],
        [
            "May not incorporate strong content-level forgery detection; primarily addresses record integrity.",
            "Blockchain integration can introduce cost/latency; solutions must be carefully engineered for usability.",
        ],
        "CertVerify includes certificate lifecycle management and a smart contract method (`revokeHash`) for revocation, while also incorporating AI-based forgery detection and cryptographic signatures to cover both visual forgery and integrity.",
    )

    # 5
    paper_block(
        "2.5 Fischinger & Boyer (2023) — DOI: 10.5281/zenodo.7326540",
        "Fischinger and Boyer, 2023",
        "They released and analyzed a dataset for document/forgery evaluation and provided an experimental baseline for classification.",
        "Dataset curation and benchmarking with baseline models, enabling reproducible evaluation of forgery detection methods.",
        [
            "Enables reproducible comparisons by providing data and baseline protocols.",
            "Encourages standardized evaluation of document-forgery models.",
        ],
        [
            "A dataset alone does not solve end-to-end credential verification; deployment concerns remain.",
            "May not be focused specifically on academic certificate templates and domain constraints.",
        ],
        "CertVerify uses CASIA v2.0 as primary training data and leverages DF2023 as supplementary data for improved robustness, then embeds the model in a full verification workflow with cryptographic and blockchain layers.",
    )

    # 6
    paper_block(
        "2.6 Alammary et al. (2019) — DOI: 10.3390/app9122400",
        "Alammary et al., 2019",
        "They surveyed blockchain-based educational credential management and compared approaches for issuing, storing, and verifying credentials.",
        "Survey of blockchain credential systems across architectures (public/private chains), storage patterns (on-chain vs off-chain), and stakeholder governance models.",
        [
            "Comprehensive overview of blockchain patterns and trade-offs for education.",
            "Highlights privacy, scalability, and interoperability challenges.",
        ],
        [
            "Survey-oriented; does not provide a concrete, evaluated system implementation.",
            "Does not address AI-based visual forgery detection for document content.",
        ],
        "CertVerify operationalizes a practical hybrid design: off-chain storage with on-chain hash anchoring, plus an AI forgery detection layer to address content-level fraud.",
    )

    # 7
    paper_block(
        "2.7 Luo et al. (2023) — IEEE TIFS 18, 3480–3495",
        "Luo et al., 2023",
        "They proposed advanced deep learning methods for image manipulation detection with improved generalization and robustness to post-processing.",
        "Forensics-oriented neural architectures with training protocols targeting generalizable manipulation cues under diverse transformations and compression.",
        [
            "Improves robustness against common post-processing operations (compression, resizing).",
            "Focuses on generalizable forensic features beyond dataset-specific artifacts.",
        ],
        [
            "Models can be computationally heavy; deployment in web systems may require optimization.",
            "Does not integrate cryptographic verification or blockchain audit mechanisms.",
        ],
        "CertVerify adopts a lightweight, deployable CNN backbone (ResNet-18) enhanced with ELA preprocessing for practical latency, and complements it with SHA-256, RSA signatures, and blockchain anchoring for end-to-end assurance.",
    )

    # 8
    paper_block(
        "2.8 Nakamura et al. (2022) — Computers & Security 112, 102510",
        "Nakamura et al., 2022",
        "They studied secure verification and management of digital documents/credentials, emphasizing threat models, authentication, and integrity guarantees.",
        "Security analysis and system design principles for document authentication; discussion of cryptographic mechanisms and practical considerations.",
        [
            "Provides rigorous security framing (threat modeling, authentication, integrity).",
            "Highlights importance of key management and secure storage.",
        ],
        [
            "Primarily security-centric; may not address computer-vision-based forgery detection.",
            "Often assumes trusted document issuance pipelines without evaluating visually forged artifacts.",
        ],
        "CertVerify integrates strong cryptography (RSA-2048, SHA-256, JWT, BCrypt) and introduces AI-based visual forgery detection for adversarially manipulated documents, producing interpretable layer-wise results and logs.",
    )

    # 9
    paper_block(
        "2.9 Salau et al. (2023) — Neural Computing and Applications 35, 8901–8920",
        "Salau et al., 2023",
        "They evaluated deep learning strategies for image forgery detection and compared preprocessing and model choices for improved classification.",
        "CNN-based classification with comparisons across preprocessing/augmentation; experimental evaluation on tampering datasets.",
        [
            "Demonstrates that preprocessing choices can significantly influence forgery detection performance.",
            "Provides comparative insights into model and training strategies.",
        ],
        [
            "Research focus is detection accuracy; does not cover issuance, signatures, or blockchain integration.",
            "Transfer to domain-specific documents (certificates) may require additional tuning.",
        ],
        "CertVerify explicitly applies ELA preprocessing (quality=90) and uses transfer learning (ImageNet-pretrained ResNet-18) tuned on CASIA v2.0, then embeds the model in a credential issuance and verification platform.",
    )

    # 10
    paper_block(
        "2.10 Zhu et al. (2022) — DOI: 10.1109/CVPR52688.2022.01453",
        "Zhu et al., 2022",
        "They presented a state-of-the-art computer vision approach for manipulation detection/localization emphasizing strong performance under challenging transformations.",
        "Advanced neural architectures for tampering detection; often includes localization heads and robust training strategies to detect subtle manipulation artifacts.",
        [
            "High accuracy and robustness on challenging benchmarks.",
            "Often supports localization of manipulated regions, improving interpretability.",
        ],
        [
            "Higher computational complexity may hinder real-time web deployment without optimization.",
            "Does not provide an integrated credential verification system with cryptographic assurance and audit trails.",
        ],
        "CertVerify prioritizes deployability and end-to-end assurance by using a compact ResNet-18 classifier with ELA for fast inference and combining it with SHA-256 integrity, RSA issuer authentication, and blockchain anchoring for auditability.",
    )

    doc.add_page_break()

    # -------------------------
    # CHAPTER 3 — Analysis of Reviewed Papers
    # -------------------------
    _add_chapter_heading(doc, "CHAPTER 3 — ANALYSIS OF REVIEWED PAPERS")
    _add_section_heading(doc, "3.1 Comparative Analysis Table")

    comparative_headers = ["Sl No", "Authors", "Methodology", "Advantages", "Disadvantages"]
    comparative_rows = [
        ("1", "Golar et al. (2025)", "Conceptual combined framework (hash + blockchain) for certificate verification", "Multi-control viewpoint; emphasizes immutability", "Not implemented; no AI forgery detection; limited workflow details"),
        ("2", "Ch et al. (2024)", "DL-based forgery detection trained with constrained data", "Demonstrates CNN utility for forgery detection", "Limited dataset scale; lacks crypto/blockchain integration"),
        ("3", "Babu et al. (2022)", "Hash-based / blockchain-stored hash verification for certificates", "Detects post-issuance tampering; automates integrity checks", "Cannot detect non-issued visual forgeries; limited non-repudiation"),
        ("4", "Vidal et al. (2020)", "Blockchain-backed certificate management with lifecycle discussion", "Highlights revocation governance; improves traceability", "Content-level forgery not addressed; chain cost/latency concerns"),
        ("5", "Fischinger & Boyer (2023)", "Dataset release and baseline benchmarking for forgery detection", "Reproducible evaluation; standardization", "Not an end-to-end system; domain mismatch possible"),
        ("6", "Alammary et al. (2019)", "Survey of blockchain credential architectures and trade-offs", "Comprehensive taxonomy; privacy/scalability insights", "Survey only; no AI forgery detection; no implementation"),
        ("7", "Luo et al. (2023)", "Robust deep learning for manipulation detection under transformations", "Better generalization; robust to compression", "Compute-heavy; no crypto or issuance workflow"),
        ("8", "Nakamura et al. (2022)", "Security-centric document verification design principles", "Threat modeling; emphasizes key management", "Limited vision-based forgery handling; assumes trusted pipeline"),
        ("9", "Salau et al. (2023)", "CNN-based forgery detection with preprocessing comparisons", "Shows preprocessing impact; good experimental analysis", "No issuance/auth/audit integration; needs domain adaptation"),
        ("10", "Zhu et al. (2022)", "SOTA vision method for manipulation detection/localization", "High performance; interpretability via localization", "High complexity; lacks end-to-end credential system"),
    ]

    table = _add_table(doc, rows=1 + len(comparative_rows), cols=len(comparative_headers))
    for j, h in enumerate(comparative_headers):
        _set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(comparative_rows, start=1):
        for j, val in enumerate(row):
            _set_cell_text(table.rows[i].cells[j], val)

    _add_paragraph(
        doc,
        "Research Gap: The reviewed literature shows that cryptographic and blockchain approaches strongly address integrity and auditability, "
        "but they do not detect visually convincing forged certificates that were never officially issued. Conversely, deep learning-based forgery "
        "detection can identify visual tampering but typically lacks issuer authentication, non-repudiation, and immutable audit trails. Moreover, "
        "many proposed frameworks remain conceptual or omit lifecycle management such as revocation. CertVerify addresses these gaps by delivering a "
        "deployed, end-to-end system that combines AI-based forgery detection (ELA + ResNet-18) with SHA-256 integrity checking, RSA-2048 digital "
        "signature validation, and Polygon blockchain anchoring with revocation support, producing transparent outcomes and verification logs."
    )

    doc.add_page_break()

    # -------------------------
    # CHAPTER 4 — Requirement Analysis
    # -------------------------
    _add_chapter_heading(doc, "CHAPTER 4 — REQUIREMENT ANALYSIS")

    _add_section_heading(doc, "4.1 Functional Requirements")
    _add_paragraph(doc, "Institution requirements:")
    inst_reqs = [
        "The system shall allow institution users to register and authenticate using email and password.",
        "The system shall enforce role-based access control such that only institutions can issue certificates.",
        "The system shall allow institutions to upload certificate files in PDF/JPG/PNG formats.",
        "The system shall compute and store SHA-256 hash and RSA-2048 digital signature for each issued certificate.",
        "The system shall generate a QR code for each certificate to enable public verification by certificate ID/URL.",
        "The system shall store certificate metadata (student name, degree, institution name, issue date) and issuance audit details.",
        "The system shall allow institutions to view a list of issued certificates and their statuses.",
        "The system shall support certificate lifecycle actions such as revocation and status updates (where applicable).",
    ]
    for idx, r in enumerate(inst_reqs, start=1):
        _add_paragraph(doc, f"{idx}. {r}")

    _add_paragraph(doc, "Verifier requirements:")
    ver_reqs = [
        "The system shall allow verifier users to register and authenticate using JWT-based login.",
        "The system shall allow verifiers to verify certificates by entering a certificate ID or by uploading a file.",
        "The system shall execute a multi-layer verification pipeline and provide layer-wise results with confidence scores.",
        "The system shall record verification history for authenticated verifiers.",
        "The system shall display final outcomes as AUTHENTIC, TAMPERED, FORGED, or INVALID.",
    ]
    for idx, r in enumerate(ver_reqs, start=1):
        _add_paragraph(doc, f"{idx}. {r}")

    _add_paragraph(doc, "Public interface requirements:")
    pub_reqs = [
        "The system shall support public verification through QR code scan leading to a verification endpoint.",
        "The system shall expose a secure REST API for verification without revealing sensitive user credentials.",
        "The system shall provide a readable verification report that can be used by employers and institutions.",
    ]
    for idx, r in enumerate(pub_reqs, start=1):
        _add_paragraph(doc, f"{idx}. {r}")

    _add_section_heading(doc, "4.2 Non-Functional Requirements")
    nfrs = [
        ("Performance", "Verification responses should be generated within practical latency limits for web use; early-exit logic should reduce average response time when forgery is detected with high confidence."),
        ("Security", "Passwords must be stored using BCrypt; API endpoints must be protected using JWT; RSA keys must be handled securely; hashes and signatures must be stored and verified without exposure of private keys."),
        ("Scalability", "The architecture should allow horizontal scaling of stateless backend services and separate storage for issued certificates if migrated from SQLite to a production database."),
        ("Reliability", "The system should maintain correctness under network failures; blockchain operations should be designed as optional/decoupled such that verification can still proceed using stored data when chain access is degraded."),
        ("Maintainability", "Code should be modular (auth, issuance, ML inference, blockchain, logging) with clear interfaces and testable components."),
        ("Portability", "The solution should run on common environments and be deployable on cloud platforms; frontend should be browser-compatible and responsive."),
    ]
    for k, v in nfrs:
        _add_paragraph(doc, f"• {k}: {v}")

    _add_section_heading(doc, "4.3 Hardware Requirements")
    hw_table = _add_table(doc, rows=5, cols=2)
    _set_cell_text(hw_table.rows[0].cells[0], "Component", bold=True)
    _set_cell_text(hw_table.rows[0].cells[1], "Specification", bold=True)
    hw_rows = [
        ("Processor", "Intel Core i5 (8th Gen) or equivalent and above"),
        ("RAM", "8 GB minimum (16 GB recommended for model training)"),
        ("Storage", "20 GB free disk space (more for datasets and model weights)"),
        ("GPU (optional)", "NVIDIA CUDA-capable GPU recommended for faster training; CPU inference supported"),
    ]
    for i, (c, s) in enumerate(hw_rows, start=1):
        _set_cell_text(hw_table.rows[i].cells[0], c)
        _set_cell_text(hw_table.rows[i].cells[1], s)

    _add_section_heading(doc, "4.4 Software Requirements")
    sw_table = _add_table(doc, rows=10, cols=2)
    _set_cell_text(sw_table.rows[0].cells[0], "Component", bold=True)
    _set_cell_text(sw_table.rows[0].cells[1], "Technology / Version", bold=True)
    sw_rows = [
        ("Programming Language", "Python 3.11 (development); Python 3.9+ supported"),
        ("Backend Framework", "FastAPI"),
        ("ORM / Database", "SQLAlchemy + SQLite"),
        ("Authentication", "JWT (token-based) + BCrypt password hashing"),
        ("Cryptography", "RSA-2048 using `cryptography`; SHA-256 using `hashlib`"),
        ("Machine Learning", "PyTorch ResNet-18; OpenCV; Pillow; ELA preprocessing"),
        ("Blockchain", "Solidity 0.8.0 smart contract; Web3.py; Polygon Amoy testnet"),
        ("Frontend", "React 18 + TypeScript; Vite; Tailwind CSS; Axios"),
        ("Deployment", "Render (backend) and Vercel (frontend)"),
    ]
    for i, (c, s) in enumerate(sw_rows, start=1):
        _set_cell_text(sw_table.rows[i].cells[0], c)
        _set_cell_text(sw_table.rows[i].cells[1], s)

    doc.add_page_break()

    # -------------------------
    # CHAPTER 5 — Design
    # -------------------------
    _add_chapter_heading(doc, "CHAPTER 5 — DESIGN")

    _add_section_heading(doc, "5.1 System Architecture")
    _add_paragraph(
        doc,
        "CertVerify follows a layered, service-oriented architecture designed for secure credential issuance and reliable verification. "
        "The Presentation Tier is a React (TypeScript) web application deployed on Vercel and provides dedicated views for home, "
        "institution issuance workflows, and verifier verification workflows. The Application Tier is implemented as a FastAPI backend "
        "deployed on Render and exposes REST endpoints for authentication, certificate issuance, verification, and audit logging. Within "
        "the backend, the Auth Module uses JWT tokens for session-less authentication and BCrypt for password storage. The Certificate "
        "Issuance Engine computes SHA-256 hashes, generates RSA-2048 digital signatures, and creates QR codes for public lookup. The AI "
        "Forgery Detection module applies Error Level Analysis preprocessing followed by ResNet-18 inference to detect visual manipulations "
        "in certificate images. The Verification Engine orchestrates the four-layer pipeline and aggregates layer-wise outcomes into a "
        "deterministic final verdict. The Data Tier uses SQLite via SQLAlchemy to persist users, certificates, and verification logs. The "
        "Blockchain Tier uses Web3.py to interact with a Solidity smart contract on Polygon Amoy, anchoring hashes and supporting revocation "
        "to provide immutable auditability."
    )

    fig_5_1_prompt = (
        "[FIGURE 5.1 — AI IMAGE PROMPT: Create a clean professional system architecture diagram for CertVerify. "
        "Show four tiers top to bottom: Presentation Tier (React frontend on Vercel with Home, Issue Certificate, "
        "Verify Certificate pages), Application Tier (FastAPI backend on Render with Auth Module JWT+BCrypt, "
        "Certificate Issuance Engine SHA-256+RSA+QR, AI Forgery Detection ELA+ResNet-18, Verification Engine, "
        "Audit Logger), Data Tier (SQLite database with Users, Certificates, Verification_Logs tables), Blockchain "
        "Tier (Polygon Amoy Testnet with Solidity Smart Contract CertRegistry). Connect with labeled arrows: HTTPS "
        "REST API, SQL queries, Web3.py RPC calls. Blue and white professional style.]"
    )
    figures.append(("FIGURE 5.1: System Architecture of CertVerify", fig_5_1_prompt))
    _add_paragraph(doc, "Figure 5.1: System Architecture of CertVerify", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=3)
    _add_paragraph(doc, fig_5_1_prompt)

    _add_section_heading(doc, "5.2 Data Flow Diagram")
    _add_paragraph(
        doc,
        "The Data Flow Diagram (DFD) captures how data moves through CertVerify during certificate issuance and verification. In issuance, "
        "an authenticated institution provides certificate metadata and uploads a file; the backend generates cryptographic artifacts "
        "(hash and signature), stores the record in the database, anchors the hash on Polygon, and returns a QR code to the institution. "
        "In verification, a verifier provides either a certificate ID or uploads a file; the backend retrieves certificate records, "
        "executes AI forgery detection, integrity verification via hash comparison, signature validation using the stored public key, "
        "and blockchain cross-check with revocation status. The final verdict and layer-wise outcomes are returned to the verifier and "
        "a verification log is recorded for auditability."
    )
    fig_5_2_prompt = (
        "[FIGURE 5.2 — AI IMAGE PROMPT: Create a professional Data Flow Diagram (DFD) / flowchart for CertVerify in blue-white theme. "
        "Include two clearly separated flows: (A) Certificate Issuance Flow and (B) Certificate Verification Flow. "
        "A: Institution user logs in -> enters certificate metadata -> uploads PDF/JPG/PNG -> backend computes SHA-256 hash -> "
        "generates RSA-2048 signature -> stores Users/Certificates in SQLite -> calls Web3.py to store hash on Polygon Amoy via "
        "storeHash(certId, hash) -> generates QR code -> returns QR + certId to institution. "
        "B: Verifier inputs certId or uploads file -> system retrieves certificate record -> Layer 1 decision diamond: "
        "ELA+ResNet-18 forgery confidence > threshold? yes -> FORGED; no -> Layer 2 decision: SHA-256 hash match? no -> TAMPERED; "
        "yes -> Layer 3 decision: RSA signature valid? no -> INVALID; yes -> Layer 4 decision: Blockchain hash matches and not revoked? "
        "no -> INVALID (or REVOKED if shown); yes -> AUTHENTIC. Include process boxes for logging to Verification_Logs and showing "
        "layer-wise report with confidence scores.]"
    )
    figures.append(("FIGURE 5.2: Data Flow Diagram (Issuance and Verification)", fig_5_2_prompt))
    _add_paragraph(doc, "Figure 5.2: Data Flow Diagram (Issuance and Verification)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=3)
    _add_paragraph(doc, fig_5_2_prompt)

    _add_section_heading(doc, "5.3 Algorithm")
    _add_paragraph(doc, "Algorithm 1: Certificate Issuance (Pseudocode)", bold=True)
    algo1 = [
        "Input: InstitutionUser u, CertificateMetadata m, CertificateFile f (PDF/JPG/PNG)",
        "Output: CertificateID certId, QRCode qr, Stored Certificate Record",
        "1. Authenticate u using JWT; verify u.role == INSTITUTION",
        "2. Validate metadata m (student_name, degree, institution_name, issue_date) and file type/size of f",
        "3. Store uploaded file securely and obtain file_path",
        "4. Compute cert_hash = SHA256(bytes(f))",
        "5. Generate RSA key pair for issuer if not already available (private_key, public_key_pem)",
        "6. Compute rsa_signature = RSA_SIGN(private_key, cert_hash)",
        "7. Create certificate record in database with m, file_path, cert_hash, rsa_signature, public_key_pem, status=ISSUED, issued_by=u.id",
        "8. Obtain certId from database insert",
        "9. Call smart contract on Polygon Amoy: tx = storeHash(certId, cert_hash)",
        "10. Store blockchain_tx_hash = tx.hash in certificate record",
        "11. Generate QR code encoding verification URL or certId",
        "12. Store qr_path and return (certId, qr) to institution",
        "13. Write issuance audit entry (optional) and respond success",
    ]
    for line in algo1:
        _add_paragraph(doc, line)

    _add_paragraph(doc, "Algorithm 2: Certificate Verification (Pseudocode)", bold=True, space_before_pt=6)
    algo2 = [
        "Input: VerificationRequest r (certId or uploaded file f), Optional VerifierUser v",
        "Output: Verdict in {AUTHENTIC, TAMPERED, FORGED, INVALID} with confidence and layer-wise report",
        "1. If v is authenticated, validate JWT; else proceed as public verification (no history linkage)",
        "2. If r includes certId, fetch certificate record C from database; else attempt to extract certId from QR (if present) or return INVALID",
        "3. If C does not exist or C.status indicates revoked/invalid, set verdict = INVALID and go to Step 10",
        "4. Layer 1 (AI Forgery Detection):",
        "   4.1 Convert input certificate to image representation (if PDF, render page to image)",
        "   4.2 Apply ELA preprocessing (quality=90) to obtain ela_image",
        "   4.3 Run ResNet-18 inference on ela_image to obtain (forgery_label, forgery_confidence)",
        "   4.4 If forgery_label == FORGED and forgery_confidence >= threshold, set verdict = FORGED and go to Step 10",
        "5. Layer 2 (SHA-256 Integrity):",
        "   5.1 Compute hash_new = SHA256(bytes(f))",
        "   5.2 If hash_new != C.cert_hash, set verdict = TAMPERED and go to Step 10",
        "6. Layer 3 (RSA Signature Validation):",
        "   6.1 Load public key from C.public_key_pem",
        "   6.2 Verify RSA signature: RSA_VERIFY(public_key, C.rsa_signature, C.cert_hash)",
        "   6.3 If verification fails, set verdict = INVALID and go to Step 10",
        "7. Layer 4 (Blockchain Cross-check):",
        "   7.1 Query smart contract getHash(C.certId) and certExists(C.certId)",
        "   7.2 If on-chain record missing or on-chain hash != C.cert_hash, set verdict = INVALID and go to Step 10",
        "   7.3 If revoked, set verdict = INVALID (or REVOKED if represented) and go to Step 10",
        "8. If all layers pass, set verdict = AUTHENTIC",
        "9. Construct response with verdict, forgery_label, forgery_confidence, and results of each layer",
        "10. Insert entry into Verification_Logs with certificate_id, result, forgery_label, forgery_confidence, ip_address, verifier_user_id, verified_at",
        "11. Return response",
    ]
    for line in algo2:
        _add_paragraph(doc, line)

    _add_section_heading(doc, "5.4 Database Design")
    _add_paragraph(
        doc,
        "CertVerify uses a relational schema with three main tables: `users`, `certificates`, and `verification_logs`. "
        "The schema captures authentication and roles, certificate issuance metadata and cryptographic artifacts, and "
        "verification audit trails. The design supports one-to-many relationships between institutions and issued certificates, "
        "and between certificates and verification events."
    )

    _add_paragraph(doc, "Table: users", bold=True)
    users_table = _add_table(doc, rows=1 + 6, cols=3)
    for j, h in enumerate(["Column", "Type", "Description"]):
        _set_cell_text(users_table.rows[0].cells[j], h, bold=True)
    users_cols = [
        ("id", "INTEGER (PK)", "Unique user identifier"),
        ("email", "VARCHAR (UNIQUE)", "User email address for login"),
        ("hashed_password", "VARCHAR", "BCrypt-hashed password"),
        ("role", "VARCHAR", "Role of user: INSTITUTION or VERIFIER"),
        ("created_at", "DATETIME", "Account creation timestamp"),
        ("(additional constraints)", "—", "JWT tokens are derived at login; password never stored in plain text"),
    ]
    for i, (c, t, d) in enumerate(users_cols, start=1):
        _set_cell_text(users_table.rows[i].cells[0], c)
        _set_cell_text(users_table.rows[i].cells[1], t)
        _set_cell_text(users_table.rows[i].cells[2], d)

    _add_paragraph(doc, "Table: certificates", bold=True, space_before_pt=6)
    certs_table = _add_table(doc, rows=1 + 14, cols=3)
    for j, h in enumerate(["Column", "Type", "Description"]):
        _set_cell_text(certs_table.rows[0].cells[j], h, bold=True)
    cert_cols = [
        ("id", "INTEGER (PK)", "Unique certificate identifier (certId)"),
        ("student_name", "VARCHAR", "Student name on certificate"),
        ("degree", "VARCHAR", "Degree/program name"),
        ("institution_name", "VARCHAR", "Issuing institution name"),
        ("issue_date", "DATE", "Certificate issue date"),
        ("file_path", "VARCHAR", "Server-side path/location of uploaded certificate"),
        ("cert_hash", "VARCHAR(64)", "SHA-256 hash of certificate file"),
        ("rsa_signature", "TEXT/BLOB", "RSA-2048 signature over cert_hash"),
        ("public_key_pem", "TEXT", "Issuer public key in PEM format"),
        ("blockchain_tx_hash", "VARCHAR", "Blockchain transaction hash storing cert_hash"),
        ("qr_path", "VARCHAR", "Path/location of generated QR code image"),
        ("status", "VARCHAR", "Lifecycle status (e.g., ISSUED / REVOKED / INVALID)"),
        ("created_at", "DATETIME", "Issuance timestamp"),
        ("issued_by_user_id", "INTEGER (FK)", "Foreign key referencing users.id (issuer)"),
    ]
    for i, (c, t, d) in enumerate(cert_cols, start=1):
        _set_cell_text(certs_table.rows[i].cells[0], c)
        _set_cell_text(certs_table.rows[i].cells[1], t)
        _set_cell_text(certs_table.rows[i].cells[2], d)

    _add_paragraph(doc, "Table: verification_logs", bold=True, space_before_pt=6)
    logs_table = _add_table(doc, rows=1 + 8, cols=3)
    for j, h in enumerate(["Column", "Type", "Description"]):
        _set_cell_text(logs_table.rows[0].cells[j], h, bold=True)
    log_cols = [
        ("id", "INTEGER (PK)", "Unique verification log identifier"),
        ("certificate_id", "INTEGER (FK)", "References certificates.id"),
        ("result", "VARCHAR", "Final outcome: AUTHENTIC/TAMPERED/FORGED/INVALID"),
        ("forgery_label", "VARCHAR", "Model output label (e.g., AUTHENTIC/FORGED)"),
        ("forgery_confidence", "FLOAT", "Model confidence score"),
        ("ip_address", "VARCHAR", "IP address of verification request"),
        ("verifier_user_id", "INTEGER (FK, NULLABLE)", "References users.id when verifier is authenticated"),
        ("verified_at", "DATETIME", "Verification timestamp"),
    ]
    for i, (c, t, d) in enumerate(log_cols, start=1):
        _set_cell_text(logs_table.rows[i].cells[0], c)
        _set_cell_text(logs_table.rows[i].cells[1], t)
        _set_cell_text(logs_table.rows[i].cells[2], d)

    er_prompt = (
        "[FIGURE 5.3 — AI IMAGE PROMPT: Create a clean ER diagram for CertVerify database. Use crow’s foot notation. "
        "Entities: USERS (id PK, email, hashed_password, role, created_at), CERTIFICATES (id PK, student_name, degree, "
        "institution_name, issue_date, file_path, cert_hash, rsa_signature, public_key_pem, blockchain_tx_hash, qr_path, "
        "status, created_at, issued_by_user_id FK), VERIFICATION_LOGS (id PK, certificate_id FK, result, forgery_label, "
        "forgery_confidence, ip_address, verifier_user_id FK nullable, verified_at). Relationships: USERS (1) issues (M) "
        "CERTIFICATES via issued_by_user_id; CERTIFICATES (1) has (M) VERIFICATION_LOGS via certificate_id; USERS (1) "
        "performs (M) VERIFICATION_LOGS via verifier_user_id (optional). Professional blue-white style, readable labels.]"
    )
    figures.append(("FIGURE 5.3: ER Diagram for CertVerify Database", er_prompt))
    _add_paragraph(doc, "Figure 5.3: ER Diagram for CertVerify Database", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=3)
    _add_paragraph(doc, er_prompt)

    doc.add_page_break()

    # -------------------------
    # CHAPTER 6 — Conclusion
    # -------------------------
    _add_chapter_heading(doc, "CHAPTER 6 — CONCLUSION")
    concl_paras = [
        (
            "This project successfully designed and developed CertVerify, a full-stack web application for issuing and verifying "
            "academic credentials with strong resistance to forgery and tampering. The platform supports institutional issuance "
            "of certificates in common digital formats, generates cryptographic hashes and RSA signatures, produces QR codes for "
            "public verification, and records verification events for auditability. Verifiers can validate credentials by ID or "
            "file upload and receive a transparent layer-wise verification report."
        ),
        (
            "The primary research contribution is the integration of computer-vision-based forgery detection with conventional "
            "cryptographic integrity and authenticity guarantees, and optional blockchain anchoring for immutable audit. Unlike "
            "hash-only or blockchain-only systems, CertVerify explicitly evaluates whether the visual content is forged and "
            "simultaneously ensures that legitimately issued certificates cannot be modified without detection. Furthermore, the "
            "design includes role-based workflows and lifecycle considerations such as revocation, aligning the implementation with "
            "practical institutional requirements."
        ),
        (
            "The ResNet-18 model trained using ELA preprocessing on the CASIA v2.0 dataset achieved 94.67% accuracy with Precision "
            "0.9564, Recall 0.9360, and F1-score 0.9461, indicating strong capability to discriminate between authentic and tampered "
            "documents under benchmark conditions. These results provide empirical support for deploying AI as the first verification "
            "layer to detect visually forged credentials and to reduce reliance on manual inspection."
        ),
        (
            "Future work includes migrating from testnet to production-grade blockchain infrastructure with robust key management and "
            "institutional governance, enhancing certificate lifecycle handling with explicit REVOKED outcomes, and retraining or "
            "fine-tuning the forgery detection model on domain-specific certificate datasets for improved real-world generalization. "
            "Additional enhancements may include bulk verification interfaces for employers, mobile-friendly verification apps "
            "integrated with camera-based QR scanning, and performance optimizations such as model quantization and caching to support "
            "high-throughput verification scenarios."
        ),
    ]
    for p in concl_paras:
        _add_paragraph(doc, p)

    doc.add_page_break()

    # -------------------------
    # SECTION 10: References (25 items, IEEE style as specified)
    # -------------------------
    _add_section_heading(doc, "REFERENCES")
    refs = [
        "[1] Golar et al. 2025 IJACECT DOI 10.65521/ijacect.v14i3s.1629",
        "[2] Ch et al. 2024 IJoSI DOI 10.6977/IJoSI.202403_8(1).0001",
        "[3] Babu et al. 2022 IJACSA 13(5)",
        "[4] Vidal et al. 2020 IEEE ICSESS DOI 10.1109/ICSESS49925.2020.9237735",
        "[5] Fischinger & Boyer 2023 IMVIP DOI 10.5281/zenodo.7326540",
        "[6] Alammary et al. 2019 Applied Sciences DOI 10.3390/app9122400",
        "[7] Luo et al. 2023 IEEE TIFS 18 3480-3495",
        "[8] Nakamura et al. 2022 Computers & Security 112 102510",
        "[9] Salau et al. 2023 Neural Computing Applications 35 8901-8920",
        "[10] Zhu et al. 2022 CVPR DOI 10.1109/CVPR52688.2022.01453",
        "[11] Shahriar et al. 2021 IEEE IRI DOI 10.1109/IRI51233.2021.00051",
        "[12] Dong et al. 2013 IEEE CISP CASIA dataset",
        "[13] He et al. 2016 CVPR ResNet",
        "[14] Dosovitskiy et al. 2021 ICLR Vision Transformer",
        "[15] Rivest et al. 1978 CACM RSA",
        "[16] Liu et al. 2022 IEEE TCSVT PSCC-Net",
        "[17] Kwon et al. 2021 WACV CAT-Net",
        "[18] Zhang et al. 2022 IEEE TPAMI MVSS-Net",
        "[19] Nakamoto 2008 Bitcoin whitepaper",
        "[20] Goodfellow et al. 2016 Deep Learning MIT Press",
        "[21] Chollet 2017 Deep Learning with Python Manning",
        "[22] Huang et al. 2017 CVPR DenseNet",
        "[23] Rossler et al. 2019 ICCV FaceForensics++",
        "[24] Swan 2015 Blockchain OReilly",
        "[25] Boneh & Shoup 2023 Applied Cryptography Stanford",
    ]
    for r in refs:
        _add_paragraph(doc, r)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH, student_fill_markers, figures


if __name__ == "__main__":
    out, markers, figs = build_report()
    print(out)
    print("STUDENT_TO_FILL:")
    for m in markers:
        print(m)
    print("FIGURES:")
    for cap, prompt in figs:
        print(cap)
        print(prompt)
